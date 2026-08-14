"""Regenerate the binary T1 fixtures (cp1252 CSV, DOCX resumes, PDF resumes).

Committed alongside the fixtures so they are reproducible artifacts, not
mystery blobs: `python tools/build_t1_binary_fixtures.py` rewrites them.
"""
from __future__ import annotations

import unicodedata
from pathlib import Path

import docx

T1 = Path(__file__).resolve().parent.parent / "goldens" / "t1"


def minimal_pdf(pages: list[list[str]]) -> bytes:
    """Tiny single-font PDF writer (PDF 1.4, Helvetica, ASCII-ish text) —
    enough to give the golden corpus real .pdf fixtures without adding a
    PDF-writing dependency (RESUME_PLAN R1)."""

    def esc(s: str) -> str:
        return s.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")

    objs: list[bytes] = []
    page_ids = [4 + 2 * i for i in range(len(pages))]
    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    objs.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objs.append(
        f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for i, lines in enumerate(pages):
        body = " ".join(f"({esc(ln)}) Tj T*" for ln in lines)
        stream = f"BT /F1 11 Tf 72 760 Td 14 TL {body} ET".encode(
            "latin-1", "replace")
        objs.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> "
            f"/Contents {page_ids[i] + 1} 0 R >>".encode())
        objs.append(b"<< /Length %d >>\nstream\n%s\nendstream"
                    % (len(stream), stream))
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for num, obj in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{num} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objs) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_at}\n%%EOF\n").encode()
    return bytes(out)


def cp1252_roster() -> None:
    # P20 The Encoding Victim: é stored as 0xE9 — utf-8 decoding fails, the
    # deterministic cp1252 fallback must recover the accents intact.
    rows = (
        "name,email,phone,current_company,title\n"
        "Renée Fontaine,renee.fontaine@example.com,+33 6 12 34 56 78,"
        "Café Lumière,Chef de Projet\n"
    )
    (T1 / "t1_cp1252.csv").write_bytes(rows.encode("cp1252"))


def _resume(name: str, lines: list[str], table: list[list[str]] | None = None) -> None:
    doc = docx.Document()
    for line in lines:
        doc.add_paragraph(line)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, cell_text in enumerate(row):
                t.rows[r].cells[c].text = cell_text
    doc.save(str(T1 / name))


def resumes() -> None:
    _resume("resume_p01.docx", [
        "Avery Stone",
        "avery.stone@example.com | +1 415 555 2671",
        "",
        "Staff Data Engineer at Marigold Data since Feb 2021",
        "Skills: Python, Airflow, SQL",
    ], table=[["Terraform", "Airflow"]])  # R4: skills grids are a top-3 pattern
    # P03: the name line is deliberately NFD-decomposed — byte-different from
    # the roster's NFC "Núñez, Carlos", visually identical.
    _resume("resume_p03.docx", [
        unicodedata.normalize("NFD", "Carlos Núñez"),
        "carlos.nunez@example.com",
        "",
        "Platform Engineer at Vertex Cloud since 2022",
        "Skills: Kubernetes, Terraform",
    ])
    _resume("resume_p07.docx", [
        "Curriculum Vitae",  # name heuristic must decline this heading
        "Ishaan Verma — ishaan.verma@example.com",
        "",
        "Analyst at Quanta Insights 2019 - 2021",
        "Consultant at Meridian Partners 03/04/2021 - 11/2021",
        "",
        "B.Tech in Data Engineering, IIT Bombay, 2018",  # R2
    ])
    _resume("resume_p18.docx", [
        "Tomas Eder | Frontend Lead",  # R5: pipe-separated contact line
        "tomas.eder@example.com",
        "https://www.linkedin.com/in/tomas-eder?utm_campaign=profile",
        "https://github.com/teder",
        "Portfolio: https://tomas.dev",
        "",
        "Frontend Lead at Pixelforge since 2023",
    ])


def pdf_resume() -> None:
    # P21 "The Portable": the real-PDF happy path (RESUME_PLAN R1), carrying
    # one block-form experience (R3) and one education line (R2) — asserted
    # by separate test functions so failures name their mechanism.
    (T1 / "resume_p21.pdf").write_bytes(minimal_pdf([[
        "Wale Adeyemi | Platform Engineer",
        "wale.adeyemi@example.com | +44 7911 123456",
        "",
        "Harmattan Cloud -- Platform Engineer",
        "Mar 2022 - Present",
        "",
        "B.Sc in Computer Science, University of Lagos, 2016",
        "",
        "Skills: Python, Terraform, Kafka",
    ]]))


if __name__ == "__main__":
    cp1252_roster()
    resumes()
    pdf_resume()
    print("T1 binary fixtures rebuilt in", T1)
