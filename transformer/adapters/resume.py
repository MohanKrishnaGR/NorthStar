"""Resume adapter (M7 stretch): PDF/DOCX text -> the shared free-text scanner.

The only new machinery is text extraction; every rule is notes_txt's, at
resume trust (0.70). Dependencies are optional (pip install .[resume]) and
imported lazily — without them a resume file is reported `skipped`, never a
crash. A scanned image-only PDF yields no text and is reported the same way.
"""
from __future__ import annotations

import re
from pathlib import Path

from ..models import Evidence, SourceRecord
from ..normalize import text
from .base import SourceResult
from .notes_txt import scan_into

SOURCE_TYPE = "resume"

# First-line-is-the-name heuristic: 2-4 words, each starting uppercase, no
# digits/@/|. Word bodies allow unicode letters so "Carlos Núñez" qualifies;
# headings like "Curriculum Vitae" match the shape and are excluded by name.
_NAME_LINE_RE = re.compile(
    r"^(?:[A-Z][^\s\d@|,;:]{1,24})(?:\s+[A-Z][^\s\d@|,;:]{1,24}){1,3}$"
)
_NOT_NAMES = {"curriculum vitae", "resume", "cv", "personal profile"}
# "Jane Doe | Senior Engineer" contact lines: the name is the first segment.
_SEGMENT_SPLIT_RE = re.compile(r"\s*[|·•]\s*")


def detect(path: Path) -> bool:
    return path.suffix.lower() in {".pdf", ".docx"}


def extract_text(path: Path) -> str:
    """The exact prose the pipeline scans for this file (docx or pdf).

    Shared by extract() and the workspace's preview endpoint, so a preview
    can never diverge from what the engine actually reads."""
    if path.suffix.lower() == ".docx":
        return _docx_text(path)
    return _pdf_text(path)


def extract(path: Path, res: SourceResult, ctx: dict) -> None:
    body = extract_text(path)
    if not body.strip():
        raise ValueError("no extractable text (scanned/image-only file?)")

    rec = SourceRecord(record_id=f"{res.source_id}#file",
                       source_id=res.source_id, source_type=SOURCE_TYPE)

    first = next((ln.strip() for ln in body.splitlines() if ln.strip()), "")
    seg = _SEGMENT_SPLIT_RE.split(first, maxsplit=1)[0].strip()
    if _NAME_LINE_RE.match(seg) and text.fold(seg) not in _NOT_NAMES:
        at = body.find(seg)
        rec.evidence.append(Evidence(
            field_path="full_name", value=text.nfc(seg), raw_value=seg,
            source_id=res.source_id, source_type=SOURCE_TYPE,
            method="regex:resume_title_name_v1", record_id=rec.record_id,
            order_index=0,
            locator={"kind": "span", "start": at, "end": at + len(seg)}))

    scan_into(rec, body, ctx)
    res.records_read = 1
    if rec.evidence:
        res.records.append(rec)


def _docx_text(path: Path) -> str:
    import docx  # lazy: optional dependency

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    # Tables (skills grids are a top-3 resume pattern). Limitation, stated:
    # table text is appended after the paragraphs rather than at its visual
    # position — span locators still ground correctly because scanning runs
    # over exactly this joined text.
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _pdf_text(path: Path) -> str:
    import pdfplumber  # lazy: optional dependency

    with pdfplumber.open(str(path)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return _clean_pdf_text(pages)


def _clean_pdf_text(pages: list[str]) -> str:
    """Bounded PDF hygiene (RESUME_PLAN R6): drop per-page repeated
    headers/footers (keep the first occurrence), heal hyphen-split words.
    Two-column reading order is deliberately NOT guessed at — that is the
    ML-extractor seat ADR-003 reserved."""
    if len(pages) > 1:
        counts: dict[str, int] = {}
        for page in pages:
            for line in set(page.splitlines()):
                counts[line] = counts.get(line, 0) + 1
        seen: set[str] = set()
        cleaned = []
        for page in pages:
            kept = []
            for line in page.splitlines():
                repeated = (line.strip() and len(line) <= 60
                            and counts.get(line, 0) == len(pages))
                if repeated:
                    if line in seen:
                        continue
                    seen.add(line)
                kept.append(line)
            cleaned.append("\n".join(kept))
        pages = cleaned
    text_all = "\n".join(pages)
    # Join "migra-\ntion" but never "co-\nFounder": lowercase on both sides.
    return re.sub(r"([a-z])-\n([a-z])", r"\1\2", text_all)
