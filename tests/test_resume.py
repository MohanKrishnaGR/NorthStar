"""M7 stretch: resume adapter (skipped cleanly when extras not installed)."""
from pathlib import Path

import pytest

pytest.importorskip("docx", reason="resume extras not installed (pip install .[resume])")

from transformer.adapters import resume
from transformer.adapters.base import run_adapter
from transformer.pipeline import run_pipeline
from transformer.projection.config import load

ROOT = Path(__file__).resolve().parent.parent
SAMPLES = ROOT / "samples"
FIXTURE = Path(__file__).resolve().parent / "fixtures" / "resume_alice.docx"
CTX = {"default_region": None, "strict": False}


def test_resume_extracts_name_contacts_skills():
    res = run_adapter(resume, FIXTURE, CTX)
    assert res.status == "ok"
    rec = res.records[0]
    by = {}
    for e in rec.evidence:
        by.setdefault(e.field_path, []).append(e)
    assert by["full_name"][0].value == "Alice Fern"
    assert by["full_name"][0].method == "regex:resume_title_name_v1"
    assert by["emails"][0].value == "alice.fern@gmail.com"
    assert by["phones"][0].value == "+14155552671"
    assert {e.value["name"] for e in by["skills"]} >= {
        "python", "postgresql", "terraform", "airflow",
    }
    assert all(e.source_type == "resume" for e in rec.evidence)


def test_resume_merges_into_alice_cluster_with_provenance():
    files = [p for p in SAMPLES.iterdir() if p.is_file()] + [FIXTURE]
    result = run_pipeline(files, load(ROOT / "configs" / "default.json"),
                          as_of=(2026, 8))
    assert len(result.profiles) == 4  # still four people — resume joined Alice
    alice = [p for p in result.profiles if p["full_name"] == "Alice Fern"][0]
    terra = [s for s in alice["skills"] if s["name"] == "terraform"]
    assert terra and terra[0]["sources"] == ["resume_alice.docx"]
    cluster = [c for c in result.report["merges"]["clusters"]
               if "resume_alice.docx#file" in c["record_ids"]][0]
    assert "ats.json#idx=0" in cluster["record_ids"]
    # Same experience entries as without the resume: it corroborates, not dupes.
    assert len(alice["experience"]) == 2


def test_garbage_docx_is_contained(tmp_path):
    bad = tmp_path / "broken.docx"
    bad.write_bytes(b"this is not a zip archive")
    res = run_adapter(resume, bad, CTX)
    assert res.status == "skipped" and res.errors


def test_docx_table_skills_are_read(tmp_path):
    import docx as docx_mod

    doc = docx_mod.Document()
    doc.add_paragraph("Skye Grid")
    doc.add_paragraph("skye.grid@example.com")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Terraform"
    t.rows[0].cells[1].text = "Kubernetes"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    res = run_adapter(resume, path, CTX)
    names = {e.value["name"] for e in res.records[0].evidence
             if e.field_path == "skills"}
    assert {"terraform", "kubernetes"} <= names


def test_contact_line_name_splits_on_pipe(tmp_path):
    import docx as docx_mod

    doc = docx_mod.Document()
    doc.add_paragraph("Zed Pipe | Staff Engineer")
    doc.add_paragraph("zed.pipe@example.com")
    path = tmp_path / "pipe.docx"
    doc.save(str(path))
    res = run_adapter(resume, path, CTX)
    names = [e for e in res.records[0].evidence if e.field_path == "full_name"]
    assert names and names[0].value == "Zed Pipe"


def test_pdf_hygiene_headers_and_hyphens(tmp_path):
    pytest.importorskip("pdfplumber")
    import sys

    sys.path.insert(0, str(ROOT / "tools"))
    from build_t1_binary_fixtures import minimal_pdf

    path = tmp_path / "two_page.pdf"
    path.write_bytes(minimal_pdf([
        ["CONFIDENTIAL - Skye Grid",
         "skye.grid@example.com",
         "Led the datacenter migra-",
         "tion program end to end."],
        ["CONFIDENTIAL - Skye Grid",
         "Second page content."],
    ]))
    from transformer.adapters.resume import _pdf_text

    text_all = _pdf_text(path)
    assert text_all.count("CONFIDENTIAL") == 1  # repeated header kept once
    assert "migration" in text_all              # hyphen-split word healed
    assert "migra-" not in text_all
